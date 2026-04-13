"""
Classroom RAG Service — DB-backed file storage + in-memory FAISS per classroom.

Files and their embeddings are stored in PostgreSQL (ClassroomFile / ClassroomChunk).
FAISS indices are built on first query and cached in memory; they are invalidated
whenever a file is uploaded or deleted.

NOTE: This replaces the old disk-based LangChain FAISS implementation.
The old helper functions (ingest_document, query_classroom_rag, delete_classroom_index)
are kept as thin compatibility shims so existing callers still work.
"""

import io
import numpy as np
import logging
from sqlalchemy.orm import Session
from fastapi import HTTPException
from db_models import ClassroomFile, ClassroomChunk
from models import HKBUEmbeddings
from core.config import (
    API_KEY, BASE_URL, FAISS_EMBEDDING_MODEL, FAISS_EMBEDDING_API_VERSION,
    MAX_UPLOAD_SIZE_BYTES, ALLOWED_MIME_TYPES
)

logger = logging.getLogger(__name__)

try:
    import faiss
    _FAISS_AVAILABLE = True
except ImportError:
    _FAISS_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    _PYMUPDF_AVAILABLE = True
except ImportError:
    _PYMUPDF_AVAILABLE = False

try:
    import docx  # python-docx
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    from sentence_transformers import SentenceTransformer as _SentenceTransformer
    _SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    _SENTENCE_TRANSFORMERS_AVAILABLE = False


# ---------------------------------------------------------------------------
# Embedding helper — HKBU primary, sentence-transformers offline fallback
# ---------------------------------------------------------------------------

_embedder = None
_using_local_fallback = False


def _get_embedder():
    """Return embedder, falling back to local sentence-transformers if HKBU fails."""
    global _embedder, _using_local_fallback
    if _embedder is None:
        try:
            candidate = HKBUEmbeddings(
                api_key=API_KEY,
                base_url=BASE_URL,
                model=FAISS_EMBEDDING_MODEL,
                api_version=FAISS_EMBEDDING_API_VERSION,
            )
            # Quick probe to confirm the API is reachable
            candidate.embed_query("test")
            _embedder = candidate
            _using_local_fallback = False
            print("✅ Classroom RAG: using HKBU embeddings")
        except Exception as e:
            print(f"⚠️ Classroom RAG: HKBU embeddings unavailable ({e})")
            if _SENTENCE_TRANSFORMERS_AVAILABLE:
                print("   📦 Falling back to local sentence-transformers (offline)")
                _embedder = _SentenceTransformer('all-MiniLM-L6-v2')
                _using_local_fallback = True
            else:
                raise RuntimeError(
                    "No embedding model available. "
                    "HKBU failed and sentence-transformers is not installed."
                )
    return _embedder


def _embed(text: str) -> list:
    """Embed a single text string, handling both HKBU and local fallback."""
    global _embedder, _using_local_fallback
    embedder = _get_embedder()
    try:
        if _using_local_fallback:
            return embedder.encode(text, convert_to_numpy=True).tolist()
        return embedder.embed_query(text)
    except Exception as e:
        # If HKBU fails mid-session (e.g. quota expires), switch to local
        if not _using_local_fallback and _SENTENCE_TRANSFORMERS_AVAILABLE:
            print(f"⚠️ Classroom RAG: HKBU embed failed ({e}), switching to local fallback")
            _embedder = _SentenceTransformer('all-MiniLM-L6-v2')
            _using_local_fallback = True
            return _embedder.encode(text, convert_to_numpy=True).tolist()
        raise


# ---------------------------------------------------------------------------
# FILE UPLOAD VALIDATION
# ---------------------------------------------------------------------------

def validate_uploaded_file(file_content_type: str) -> None:
    """Validate file MIME type against whitelist.
    
    Args:
        file_content_type: MIME type of uploaded file
        
    Raises:
        HTTPException(400): If MIME type not in whitelist
    """
    if file_content_type not in ALLOWED_MIME_TYPES:
        allowed = ", ".join(sorted(ALLOWED_MIME_TYPES))
        logger.warning(f"File upload rejected: unsupported MIME type '{file_content_type}'")
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {file_content_type}. Allowed: {allowed}"
        )
    logger.debug(f"File MIME type valid: {file_content_type}")


async def check_file_size(file_bytes: bytes) -> None:
    """Validate file size doesn't exceed limit.
    
    Args:
        file_bytes: Uploaded file content
        
    Raises:
        HTTPException(413): If file exceeds MAX_UPLOAD_SIZE_BYTES
    """
    file_size_mb = len(file_bytes) / (1024 * 1024)
    if len(file_bytes) > MAX_UPLOAD_SIZE_BYTES:
        logger.warning(f"File upload rejected: size {file_size_mb:.1f}MB exceeds {MAX_UPLOAD_SIZE_BYTES/(1024*1024):.0f}MB limit")
        raise HTTPException(
            status_code=413,
            detail=f"File too large: {file_size_mb:.1f}MB. Max allowed: {MAX_UPLOAD_SIZE_BYTES/(1024*1024):.0f}MB"
        )
    logger.debug(f"File size valid: {file_size_mb:.1f}MB")


# ---------------------------------------------------------------------------
# In-memory FAISS index cache  { classroom_id: (index, [chunk_text, ...]) }
# ---------------------------------------------------------------------------

_index_cache: dict = {}


# ---------------------------------------------------------------------------
# Text extraction with page tracking
# ---------------------------------------------------------------------------

def _extract_text_with_pages(filename: str, data: bytes) -> list:
    """
    Extract text from file and return list of tuples: [(text, page_num), ...]
    page_num is 1-indexed for PDFs, 0 for other file types.
    """
    fname = filename.lower()
    if fname.endswith(".pdf"):
        if _PYMUPDF_AVAILABLE:
            doc = fitz.open(stream=data, filetype="pdf")
            results = []
            for page_num, page in enumerate(doc, start=1):
                text = page.get_text()
                if text.strip():
                    results.append((text, page_num))
            return results
        # Fallback: try to decode as text
        text = data.decode("utf-8", errors="ignore")
        return [(text, 0)]
    elif fname.endswith(".docx"):
        if _DOCX_AVAILABLE:
            document = docx.Document(io.BytesIO(data))
            text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
            return [(text, 0)]
        text = data.decode("utf-8", errors="ignore")
        return [(text, 0)]
    else:
        text = data.decode("utf-8", errors="ignore")
        return [(text, 0)]


def _extract_text(filename: str, data: bytes) -> str:
    """Legacy function for backwards compatibility."""
    fname = filename.lower()
    if fname.endswith(".pdf"):
        if _PYMUPDF_AVAILABLE:
            doc = fitz.open(stream=data, filetype="pdf")
            return "\n".join(page.get_text() for page in doc)
        # Fallback: try to decode as text
        return data.decode("utf-8", errors="ignore")
    elif fname.endswith(".docx"):
        if _DOCX_AVAILABLE:
            document = docx.Document(io.BytesIO(data))
            return "\n".join(p.text for p in document.paragraphs if p.text.strip())
        return data.decode("utf-8", errors="ignore")
    else:
        return data.decode("utf-8", errors="ignore")


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

def _split_chunks(text: str, chunk_size: int = 500, overlap: int = 50) -> list:
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk = " ".join(words[i: i + chunk_size])
        if chunk.strip():
            chunks.append(chunk)
        i += chunk_size - overlap
    return chunks


# ---------------------------------------------------------------------------
# Upload & Index  (new DB-backed approach)
# ---------------------------------------------------------------------------

def upload_and_index(
    classroom_id: int,
    filename: str,
    mime_type: str,
    raw_bytes: bytes,
    uploaded_by: int,
    db: Session,
    section_id: int = None,
) -> int:
    """
    1. Persist raw file bytes → classroom_files
    2. Extract text with page tracking, chunk, embed
    3. Persist chunks + embeddings + page numbers → classroom_chunks
    4. Bust in-memory FAISS cache for this classroom
    Returns: file_id
    """
    MAX_SIZE = 20 * 1024 * 1024  # 20 MB
    if len(raw_bytes) > MAX_SIZE:
        raise ValueError("File too large (max 20 MB)")

    # Persist file
    db_file = ClassroomFile(
        classroom_id=classroom_id,
        filename=filename,
        mime_type=mime_type,
        file_data=raw_bytes,
        uploaded_by=uploaded_by,
        section_id=section_id,
    )
    db.add(db_file)
    db.flush()  # populate db_file.id before inserting chunks

    # Extract text with page tracking → chunk → embed
    pages_data = _extract_text_with_pages(filename, raw_bytes)
    if not pages_data:
        db.commit()
        return db_file.id

    # Flatten all text for chunking, but track original page numbers
    all_chunks_with_pages = []
    embedder = _get_embedder()

    # Process each page separately to maintain page numbers
    for page_text, page_num in pages_data:
        if not page_text.strip():
            continue
        chunks = _split_chunks(page_text)
        for chunk in chunks:
            all_chunks_with_pages.append((chunk, page_num))

    if not all_chunks_with_pages:
        db.commit()
        return db_file.id

    # Batch embed (16 per call)
    all_embeddings = []
    batch_size = 16
    chunk_texts = [c[0] for c in all_chunks_with_pages]
    
    try:
        for i in range(0, len(chunk_texts), batch_size):
            batch = chunk_texts[i: i + batch_size]
            embeddings = embedder.embed_documents(batch)
            all_embeddings.extend(embeddings)
        print(f"✓ Successfully embedded {len(chunk_texts)} chunks for file {filename}")
    except Exception as e:
        # Rollback file if embedding fails
        print(f"❌ Embedding failed: {str(e)}")
        db.rollback()
        raise ValueError(f"Failed to embed document chunks: {str(e)}")

    # Persist chunks with page numbers
    for (chunk_text, page_num), emb in zip(all_chunks_with_pages, all_embeddings):
        emb_bytes = np.array(emb, dtype=np.float32).tobytes()
        db.add(ClassroomChunk(
            file_id=db_file.id,
            classroom_id=classroom_id,
            chunk_text=chunk_text,
            embedding=emb_bytes,
            page_number=page_num,
        ))

    db.commit()
    _index_cache.pop(classroom_id, None)  # bust cache
    print(f"✓ Successfully indexed {len(all_embeddings)} chunks for classroom {classroom_id}")
    return db_file.id


# ---------------------------------------------------------------------------
# Delete file
# ---------------------------------------------------------------------------

def delete_classroom_file(file_id: int, classroom_id: int, db: Session) -> None:
    """
    Delete file row — classroom_chunks rows are CASCADE-deleted by the FK.
    Busts the in-memory FAISS cache.
    """
    db.query(ClassroomFile).filter(
        ClassroomFile.id == file_id,
        ClassroomFile.classroom_id == classroom_id,
    ).delete(synchronize_session=False)
    db.commit()
    _index_cache.pop(classroom_id, None)


# ---------------------------------------------------------------------------
# Build / retrieve in-memory FAISS index
# ---------------------------------------------------------------------------

def _get_or_build_index(classroom_id: int, db: Session):
    """
    Returns (faiss.IndexFlatIP, [(chunk_text, file_id, filename, mime_type, page_num), ...]) for the classroom.
    Builds from Postgres when not cached.
    Returns (None, []) when classroom has no documents.
    """
    if classroom_id in _index_cache:
        return _index_cache[classroom_id]

    rows = (
        db.query(ClassroomChunk)
        .filter(ClassroomChunk.classroom_id == classroom_id)
        .all()
    )

    if not rows or not _FAISS_AVAILABLE:
        return None, []

    embeddings = np.array(
        [np.frombuffer(r.embedding, dtype=np.float32) for r in rows],
        dtype=np.float32,
    )
    faiss.normalize_L2(embeddings)

    dim = embeddings.shape[1]
    index = faiss.IndexFlatIP(dim)
    index.add(embeddings)

    # Return chunk metadata as tuples: (text, file_id, filename, mime_type, page_num)
    chunk_metadata = []
    for r in rows:
        filename = "unknown"
        mime_type = "text/plain"
        if r.file_id:
            file_obj = db.query(ClassroomFile).filter(ClassroomFile.id == r.file_id).first()
            if file_obj:
                filename = file_obj.filename
                mime_type = file_obj.mime_type
        
        # Handle missing page_number column gracefully
        page_num = getattr(r, 'page_number', None) or 1
        
        chunk_metadata.append((r.chunk_text, r.file_id, filename, mime_type, page_num))
    
    _index_cache[classroom_id] = (index, chunk_metadata)
    return index, chunk_metadata


# ---------------------------------------------------------------------------
# Search (RAG query) — returns chunks with file metadata
# ---------------------------------------------------------------------------

def get_chunks_for_files(
    classroom_id: int,
    file_ids: list,
    db: Session,
    max_chars: int = 8000,
) -> list:
    """
    Return all chunks belonging to the given file_ids (no embedding search).
    Used when the teacher explicitly selects which files to base the quiz on.
    Returns list of dicts with same shape as search_classroom_context.
    """
    rows = (
        db.query(ClassroomChunk)
        .filter(
            ClassroomChunk.classroom_id == classroom_id,
            ClassroomChunk.file_id.in_(file_ids),
        )
        .all()
    )
    results = []
    total_chars = 0
    for r in rows:
        if total_chars >= max_chars:
            break
        file_obj = db.query(ClassroomFile).filter(ClassroomFile.id == r.file_id).first()
        filename = file_obj.filename if file_obj else "unknown"
        mime_type = file_obj.mime_type if file_obj else "text/plain"
        page_num = getattr(r, "page_number", None) or 1
        results.append({
            "text": r.chunk_text,
            "file_id": r.file_id,
            "filename": filename,
            "mime_type": mime_type,
            "page_number": page_num,
        })
        total_chars += len(r.chunk_text)
    return results


def search_classroom_context(
    classroom_id: int,
    query: str,
    db: Session,
    top_k: int = 5,
) -> list:
    """
    Returns list of dicts: [{"text": chunk_text, "file_id": id, "filename": name, "mime_type": type, "page_number": num}, ...]
    Returns [] when classroom has no uploaded documents or FAISS unavailable.
    """
    index, chunk_metadata = _get_or_build_index(classroom_id, db)
    if index is None:
        print(f"⚠️ No index found for classroom {classroom_id}")
        return []

    query_emb = np.array([_embed(query)], dtype=np.float32)
    faiss.normalize_L2(query_emb)

    scores, ids = index.search(query_emb, top_k)
    results = []
    # Lower threshold to 0.1 for better matching
    threshold = 0.1
    for score, idx in zip(scores[0], ids[0]):
        if idx != -1 and score > threshold:
            chunk_text, file_id, filename, mime_type, page_num = chunk_metadata[idx]
            results.append({
                "text": chunk_text,
                "file_id": file_id,
                "filename": filename,
                "mime_type": mime_type,
                "page_number": page_num,
            })
    
    print(f"✓ RAG search for classroom {classroom_id}: query='{query}' → {len(results)} chunks found (threshold={threshold}, raw_scores={list(scores[0][:3])})")
    return results


# ---------------------------------------------------------------------------
# Legacy compatibility shims (kept so routers/classroom.py still compiles)
# ---------------------------------------------------------------------------

def ingest_document(
    classroom_id: int,
    file_path: str,
    filename: str,
    uploaded_by: int,   # new
    db: Session,        # new
) -> int:
    """Reads file from disk and delegates to upload_and_index with user and db."""
    import os
    with open(file_path, "rb") as fh:
        raw_bytes = fh.read()
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "bin"
    mime_map = {"pdf": "application/pdf", "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "txt": "text/plain", "md": "text/markdown"}
    mime_type = mime_map.get(ext, "application/octet-stream")
    chunk_count = upload_and_index(
        classroom_id=classroom_id,
        filename=filename,
        mime_type=mime_type,
        raw_bytes=raw_bytes,
        uploaded_by=uploaded_by,
        db=db,
    )
    return chunk_count


def query_classroom_rag(classroom_id: int, question: str, k: int = 5) -> list:
    """Legacy shim — searches classroom context and returns dicts with text and metadata."""
    from database import SessionLocal
    db = SessionLocal()
    try:
        chunks = search_classroom_context(classroom_id, question, db, top_k=k)
        return [{"page_content": c["text"], "file_id": c.get("file_id"), "filename": c.get("filename"), "page_number": c.get("page_number", 1)} for c in chunks]
    finally:
        db.close()


def delete_classroom_index(classroom_id: int) -> None:
    """Legacy shim — clears both DB chunks and the in-memory cache."""
    from database import SessionLocal
    db = SessionLocal()
    try:
        db.query(ClassroomChunk).filter(ClassroomChunk.classroom_id == classroom_id).delete()
        db.query(ClassroomFile).filter(ClassroomFile.classroom_id == classroom_id).delete()
        db.commit()
    finally:
        db.close()
    _index_cache.pop(classroom_id, None)
